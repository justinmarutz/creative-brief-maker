import io
import logging
import re
from datetime import date

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

log = logging.getLogger(__name__)


def _xml_escape(s):
    """Escape text for safe insertion into ReportLab XML / DOCX paragraphs."""
    return (str(s)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _safe_url(url):
    """Strip characters that would break an XML href attribute."""
    return re.sub(r'["\'\s<>]', '', str(url))


def today_mmddyy():
    return date.today().strftime("%m%d%y")


ORIENTATION_DIMENSIONS = {
    "vertical":   "9:16 vertical - follow safe zone guides for 1:1 and 4:5 when positioning captions due to Meta crops.",
    "square":     "1:1 square - keep all text and action within centre 80% of frame.",
    "horizontal": "16:9 horizontal - standard widescreen framing.",
}


def build_filename(concept, ext="pdf"):
    d = concept.get("submission_date") or today_mmddyy()
    # Strip ALL non-alphanumeric chars from brand to prevent path traversal
    brand = re.sub(r"[^a-zA-Z0-9]", "", concept.get("brand", "Selffinder")) or "Selffinder"
    title = concept.get("title", "Brief")
    slug = re.sub(r"['\"""''‘’“”…]", "", title)
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", slug).strip("_")
    slug = re.sub(r"_+", "_", slug) or "Brief"
    return f"{d}___{brand}___Video_Creative_Brief___{slug}.{ext}"


# ── PDF ───────────────────────────────────────────────────────────────────────

def generate_pdf_bytes(concept):
    ORANGE = colors.HexColor("#E8521A")
    GRAY = colors.HexColor("#CCCCCC")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.85 * inch, rightMargin=0.85 * inch,
        topMargin=0.6 * inch, bottomMargin=0.85 * inch,
    )

    logo_s = ParagraphStyle("logo", fontName="Helvetica-Bold", fontSize=13, textColor=ORANGE)
    title_s = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=22, leading=26, spaceAfter=4)
    id_s   = ParagraphStyle("id",    fontName="Helvetica",      fontSize=10, leading=13, spaceAfter=10,
                            textColor=colors.HexColor("#888888"))
    h2_s   = ParagraphStyle("h2",    fontName="Helvetica-Bold", fontSize=14, leading=17, spaceBefore=10, spaceAfter=6)
    fi_s   = ParagraphStyle("fi",    fontName="Helvetica",      fontSize=10, leading=14, spaceAfter=3)
    sc_s   = ParagraphStyle("sc",    fontName="Helvetica",      fontSize=10, leading=14, spaceAfter=5)
    st_s   = ParagraphStyle("st",    fontName="Helvetica-Oblique", fontSize=10, leading=14, spaceAfter=5,
                            textColor=colors.HexColor("#555555"))
    bu_s   = ParagraphStyle("bu",    fontName="Helvetica",      fontSize=10, leading=14, spaceAfter=3, leftIndent=18)

    def field(label, value):
        return Paragraph(f"<b>{label}</b> {_xml_escape(value)}", fi_s)

    def hr_gray():
        return HRFlowable(width="100%", thickness=0.5, color=GRAY, spaceAfter=10, spaceBefore=4)

    brand = _xml_escape(concept.get("brand") or "Self Finder")
    title = _xml_escape(concept.get("title") or "Brief")

    concept_id = _xml_escape(concept.get("concept_id") or "")

    story = []
    story.append(Paragraph("&#169; adtechnacity", logo_s))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ORANGE, spaceAfter=14, spaceBefore=6))
    story.append(Paragraph(f"Creative Brief: {brand} - {title}", title_s))
    if concept_id:
        story.append(Paragraph(f"Brief ID: {concept_id}", id_s))

    # Section 1
    story.append(Paragraph("1. Project Overview", h2_s))
    orientation = concept.get("orientation", "vertical")
    dimensions  = ORIENTATION_DIMENSIONS.get(orientation, ORIENTATION_DIMENSIONS["vertical"])

    story.append(field("Format:",     "UGC video"))
    story.append(field("Platform:",   concept.get("platform") or "Meta"))
    story.append(field("Duration:",   concept.get("duration") or "—"))
    story.append(field("Dimensions:", dimensions))

    if concept.get("format_vibe"):
        story.append(Spacer(1, 4))
        story.append(field("Concept:", concept["format_vibe"]))

    story.append(Spacer(1, 4))
    insp = (concept.get("inspiration") or "").strip()
    ref  = (concept.get("reference_link") or "").strip()

    if insp.lower() in ("", "none", "none provided") and not ref:
        story.append(field("Inspiration video:", "None provided"))
    elif re.match(r"https?://", insp):
        safe_url = _safe_url(insp)
        story.append(Paragraph(
            f'<b>Inspiration video:</b> <a href="{safe_url}" color="#1a73e8">{_xml_escape(insp)}</a>',
            fi_s,
        ))
    else:
        story.append(field("Inspiration video:", insp or "None provided"))

    if ref:
        safe_ref = _safe_url(ref)
        story.append(Paragraph(
            f'<b>Reference link:</b> <a href="{safe_ref}" color="#1a73e8">{_xml_escape(ref)}</a>',
            fi_s,
        ))

    story.append(Spacer(1, 8))
    story.append(hr_gray())

    # Section 2 — Script
    story.append(Paragraph("2. Script", h2_s))
    script = concept.get("script") or ""
    for line in script.splitlines():
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
        elif re.match(r"^\[.*\]$", line):
            story.append(Paragraph(_xml_escape(line), st_s))
        else:
            story.append(Paragraph(_xml_escape(line), sc_s))

    story.append(Spacer(1, 8))
    story.append(hr_gray())

    # Section 3
    story.append(Paragraph("3. Production &amp; Performance Notes", h2_s))
    if concept.get("talent"):
        story.append(field("Talent:", concept["talent"]))
    story.append(field("Setting:", concept.get("setting") or "—"))
    story.append(Spacer(1, 4))
    story.append(field("Audio:", "No background music. No sound effects."))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Performance direction:</b>", fi_s))

    # format_vibe is shown in Section 1 as Concept: — don't repeat it here
    perf_bullets = [
        "No ring lights — phone front camera, natural available light only",
        "No mic — phone audio. If it sounds studio-clean, it's wrong",
        "Deliver like you're thinking of it in real time — not reciting from a script",
    ]
    for b in perf_bullets:
        story.append(Paragraph(f"•  {b}", bu_s))

    story.append(Spacer(1, 4))
    story.append(field("Editing:", "Straight cuts only. No fancy transitions."))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _add_hr(doc, color_hex, thickness_eighths=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_field(doc, label, value, size_pt=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.size = Pt(size_pt)
    v = p.add_run(str(value))
    v.font.size = Pt(size_pt)
    return p


def _add_hyperlink_field(doc, label, url):
    """Bold label + clickable blue hyperlink on the same paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label + " ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    # Add hyperlink via relationship
    part = doc.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    link_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Blue underline style
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1a73e8")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")  # 10pt in half-points
    rPr.append(color)
    rPr.append(u)
    rPr.append(sz)
    link_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = url
    link_run.append(t)
    hyperlink.append(link_run)
    p._p.append(hyperlink)
    return p


def generate_docx_bytes(concept):
    ORANGE    = RGBColor(0xE8, 0x52, 0x1A)
    GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.85)

    # Logo
    logo_p = doc.add_paragraph()
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(2)
    r = logo_p.add_run("© adtechnacity")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = ORANGE

    _add_hr(doc, "E8521A", 12)

    # Title
    brand = concept.get("brand", "Self Finder")
    title = concept.get("title", "Brief")
    concept_id = concept.get("concept_id", "")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after = Pt(4)
    tr = tp.add_run(f"Creative Brief: {brand} - {title}")
    tr.bold = True
    tr.font.size = Pt(22)

    if concept_id:
        id_p = doc.add_paragraph()
        id_p.paragraph_format.space_before = Pt(0)
        id_p.paragraph_format.space_after = Pt(10)
        id_r = id_p.add_run(f"Brief ID: {concept_id}")
        id_r.font.size = Pt(10)
        id_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Section 1
    def section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)

    section_heading("1. Project Overview")
    orientation = concept.get("orientation", "vertical")
    dimensions = ORIENTATION_DIMENSIONS.get(orientation, ORIENTATION_DIMENSIONS["vertical"])

    _add_field(doc, "Format:", "UGC video")
    _add_field(doc, "Platform:", concept.get("platform", "Meta"))
    _add_field(doc, "Duration:", concept.get("duration") or "—")
    _add_field(doc, "Dimensions:", dimensions)
    if concept.get("format_vibe"):
        _add_field(doc, "Concept:", concept["format_vibe"])
    insp = (concept.get("inspiration") or "").strip()
    _add_field(doc, "Inspiration video:", insp or "None provided")

    ref = (concept.get("reference_link") or "").strip()
    if ref:
        _add_hyperlink_field(doc, "Reference link:", ref)

    _add_hr(doc, "CCCCCC", 4)

    # Section 2 — Script
    section_heading("2. Script")
    script = concept.get("script", "")
    for line in script.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(line)
        r.font.size = Pt(10)
        if re.match(r"^\[.*\]$", line):
            r.italic = True
            r.font.color.rgb = GRAY_TEXT

    _add_hr(doc, "CCCCCC", 4)

    # Section 3
    section_heading("3. Production & Performance Notes")
    if concept.get("talent"):
        _add_field(doc, "Talent:", concept["talent"])
    _add_field(doc, "Setting:", concept.get("setting") or "—")
    _add_field(doc, "Audio:", "No background music. No sound effects.")

    pd_p = doc.add_paragraph()
    pd_p.paragraph_format.space_before = Pt(4)
    pd_p.paragraph_format.space_after = Pt(3)
    r = pd_p.add_run("Performance direction:")
    r.bold = True
    r.font.size = Pt(10)

    perf_bullets = [
        "No ring lights — phone front camera, natural available light only",
        "No mic — phone audio. If it sounds studio-clean, it's wrong",
        "Deliver like you're thinking of it in real time — not reciting from a script",
    ]
    # format_vibe shown in Section 1 as Concept: — not repeated here
    for b in perf_bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(3)
        r = bp.add_run(b)
        r.font.size = Pt(10)

    _add_field(doc, "Editing:", "Straight cuts only. No fancy transitions.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
