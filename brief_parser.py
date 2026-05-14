import datetime
import io
import re
from datetime import date


LABEL_MAP = [
    (r"^script\s*$", "script"),
    (r"^duration\s*$", "duration"),
    (r"^format\s*/\s*vibe\s*$|^format\s*$", "format_vibe"),
    (r"^inspiration\s*$", "inspiration"),
    (r"^setting\s*$", "setting"),
    (r"^talent\s*$", "talent"),
    (r"^angle\s*$", "angle"),
    (r"^on.?screen text.*$", "on_screen_text"),
    (r"^primary text.*$|^ad copy.*$", "ad_copy"),
    (r"^brand\s*$", "brand"),
    (r"^platform\s*$", "platform"),
    (r"^submission date\s*$", "submission_date"),
    (r"^campaign goal\s*$", "campaign_goal"),
]


def extract_text_from_docx(file_bytes):
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(c.text for c in row.cells))
    return "\n".join(lines)


def _match_label(line):
    clean = line.strip().rstrip(":").strip().lower()
    for pattern, key in LABEL_MAP:
        if re.match(pattern, clean, re.IGNORECASE):
            return key
    return None


def _parse_concept_block(block, index):
    today = date.today().strftime("%m%d%y")
    concept = {
        "index": index,
        "title": "",
        "brand": "SelfFinder",
        "platform": "Meta",
        "duration": "",
        "setting": "",
        "talent": "",
        "inspiration": "",
        "reference_link": "",
        "script": "",
        "format_vibe": "",
        "submission_date": today,
        "missing": [],
    }

    lines = block.strip().splitlines()
    if not lines:
        return None

    # Extract title from first line
    first = lines[0].strip()
    quoted = re.search(r'[“”„"](.*?)[“”„"]', first)
    if quoted:
        concept["title"] = quoted.group(1).strip()
    elif ":" in first:
        concept["title"] = first.split(":", 1)[1].strip().strip('"').strip("“”")
    else:
        concept["title"] = first.strip()

    # Scan lines for label→value blocks
    current_key = None
    current_lines = []

    def flush(key, val_lines):
        if not key:
            return
        # Ensure stage directions on the same line as dialogue get a line break
        joined = []
        for l in val_lines:
            if l.strip():
                # Insert newline when a closing ] is immediately followed by a quote
                l = re.sub(r'(\])\s*(["""])', r'\1\n\2', l)
                joined.append(l)
        value = "\n".join(joined).strip()
        # Only set if not already set by a more specific field
        if key in concept and concept[key]:
            return
        if key in concept:
            concept[key] = value
        elif key == "angle":
            if not concept.get("setting"):
                concept["setting"] = value
        elif key == "format_vibe":
            concept["format_vibe"] = value

    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            if current_key:
                current_lines.append("")
            continue

        matched = _match_label(stripped)
        if matched:
            flush(current_key, current_lines)
            current_key = matched
            current_lines = []
            # Inline value after colon on same line
            if ":" in stripped:
                inline = stripped.split(":", 1)[1].strip()
                if inline:
                    current_lines.append(inline)
        else:
            if current_key:
                current_lines.append(stripped)

    flush(current_key, current_lines)

    # Split inspiration into description + reference_link when both are present
    # Handles formats like: "myIQ 'are you smarter...' — https://..."
    if concept["inspiration"] and not concept["reference_link"]:
        url_match = re.search(r'(https?://\S+)', concept["inspiration"])
        if url_match:
            url = url_match.group(1).rstrip(")")  # strip trailing parens
            desc = concept["inspiration"][:url_match.start()].rstrip(" —–-").strip()
            concept["reference_link"] = url
            if desc:
                concept["inspiration"] = desc

    # Normalize duration em-dash
    if concept["duration"]:
        concept["duration"] = concept["duration"].replace("-", "–")

    # Mark missing fields
    for f in ["duration", "setting", "inspiration"]:
        if not concept.get(f):
            concept["missing"].append(f)

    return concept


ANGLE_TALENT = {
    "partner": "Woman, 28–42. Looks like she's been in a relationship — relatable, not a lifestyle influencer.",
    "self":    "Flexible — any creator who can deliver with genuine self-awareness energy.",
    "general": "Flexible — any creator who reads as credible and grounded.",
}

ANGLE_SETTING = {
    "partner": "Realistic living room on couch, or in car — anywhere casual.",
    "self":    "At home — bedroom, kitchen, or couch. Natural light.",
    "general": "At home or in car. Natural light, phone-grade.",
}


def _extract_trailing_metadata(text):
    """
    Parse the trailing tab-separated table that has per-script triplets:
      Angle\t<value>
      Duration\t<value>
      Format/Vibe\t<value>
    Returns a list of dicts in script order.
    """
    groups = []
    current = {}
    for line in text.splitlines():
        if "\t" not in line:
            continue
        key, _, val = line.partition("\t")
        key = key.strip().lower().rstrip(":")
        val = val.strip()
        if not val:
            continue
        if key == "angle":
            if current:
                groups.append(current)
            current = {"angle": val.lower()}
        elif key == "duration" and current:
            current["duration"] = val.replace("-", "–")
        elif re.match(r"format\s*/\s*vibe|format$", key) and current:
            current["format_vibe"] = val
    if current:
        groups.append(current)
    return groups


def _extract_global_metadata(text):
    """Pull submission-level fields that apply to all concepts."""
    global_data = {}
    for line in text.splitlines():
        if "\t" in line:
            key, _, val = line.partition("\t")
            key = key.strip().lower().rstrip(":")
            val = val.strip()
            if key == "product" and val:
                global_data["brand_note"] = val
            elif key == "platform" and val:
                global_data["platform"] = val
    # Detect submission date from header (e.g. "April 8, 2026")
    date_match = re.search(r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})", text)
    if date_match:
        try:
            dt = datetime.datetime.strptime(date_match.group(0).replace(",", ""), "%B %d %Y")
            global_data["submission_date"] = dt.strftime("%m%d%y")
        except ValueError:
            pass
    return global_data


# ── VIDEO N: format parser ────────────────────────────────────────────────────
# Handles submissions like SelfFinder_051126 where the structure is:
#   VIDEO 1: Title
#   Format: X Reference: Y Cast: Z   ← all inline on one line
#   V1A — Sub-variant title
#   SCRIPT:
#   ...dialogue...
#   Primary text (Meta):
#   ...
#   VIDEO 2: Title
#   ...

# Field names that can appear inline in this format (longest first to prevent
# prefix mis-matches — e.g. "Cast direction" must beat "Cast")
_VID_FIELDS = sorted([
    "Cast direction for Erik", "Cast direction",
    "Why this first", "Why this beats", "Why this",
    "Primary text (Meta)", "Primary text",
    "Persistent overlay", "Production",
    "Reference", "Headline", "Format",
    "Audio", "Cast", "CTA", "End card",
], key=len, reverse=True)

_VID_INLINE_RE = re.compile(
    r"(?<!\w)(?:" + "|".join(re.escape(f) for f in _VID_FIELDS) + r")\s*:",
    re.IGNORECASE,
)


def _vid_get(text, field_name):
    """Return the value of `field_name` from inline-concatenated field text."""
    matches = list(_VID_INLINE_RE.finditer(text))
    for i, m in enumerate(matches):
        label = m.group(0).rstrip(":").strip().lower()
        if label == field_name.lower():
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            return text[start:end].strip()
    return ""


def _vid_duration(text):
    """Extract a duration string like '30–45 sec' from free text."""
    m = re.search(r"(\d+\s*[–—-]\s*\d+\s*(?:sec|s\b|min)|\d+\s*(?:sec|s\b|min))",
                  text, re.IGNORECASE)
    return m.group(1).replace("-", "–").strip() if m else ""


def _build_vid_concept(block, idx, today, parent_meta=""):
    """Build one concept dict from a VIDEO block or sub-variant block."""
    concept = {
        "index": idx, "title": "", "brand": "SelfFinder", "platform": "Meta",
        "duration": "", "setting": "", "talent": "", "inspiration": "",
        "reference_link": "", "script": "", "format_vibe": "",
        "submission_date": today, "missing": [],
    }

    lines = block.strip().splitlines()
    if not lines:
        return None

    # ── Title from first line ────────────────────────────────────────────
    first = lines[0].strip()
    m = re.match(r"^VIDEO\s+\d+\s*:\s*(.+)", first)
    concept["title"] = m.group(1).strip() if m else first

    # ── Inline fields — search block + shared parent metadata ────────────
    search = parent_meta + "\n" + block

    fmt = _vid_get(search, "Format")
    if fmt:
        concept["format_vibe"] = fmt
        concept["duration"] = _vid_duration(fmt)

    cast = _vid_get(search, "Cast")
    if cast:
        concept["talent"] = cast

    ref = _vid_get(search, "Reference")
    if ref:
        url_m = re.search(r"https?://\S+", ref)
        if url_m:
            concept["reference_link"] = url_m.group(0).rstrip(")")
            desc = ref[: url_m.start()].rstrip(" —–-").strip()
            if desc:
                concept["inspiration"] = desc
        else:
            concept["inspiration"] = ref

    # ── Script ─────────────────────────────────────────────────────────
    # Talking-head / UGC: look for SCRIPT: label
    script_m = re.search(
        r"(?m)^SCRIPT\s*:\s*\n(.*?)(?=\nPrimary text|\nEnd card|\nCast direction|\nHeadline|\Z)",
        block, re.IGNORECASE | re.DOTALL,
    )
    if script_m:
        concept["script"] = script_m.group(1).strip()
    else:
        # Compilation format: grab body lines between metadata and Primary text
        skip_re = re.compile(
            r"^(?:Format|Reference|Production|Cast|Why this|Audio)\s*:", re.IGNORECASE
        )
        stop_re = re.compile(
            r"^(?:Primary text|End card|Headline)\b", re.IGNORECASE
        )
        body, past_meta = [], False
        for line in lines[1:]:
            ls = line.strip()
            if not ls:
                continue
            if not past_meta and skip_re.match(ls):
                continue
            if stop_re.match(ls):
                break
            past_meta = True
            body.append(ls)
        if body:
            concept["script"] = "\n".join(body)

    # ── Setting from stage directions (first non-dialogue script line) ───
    if concept["script"] and not concept["setting"]:
        skip_stage = re.compile(
            r"^(?:Clip \d+|Sequence|Persistent|Audio|Primary|V\d+[A-Z])", re.IGNORECASE
        )
        for line in concept["script"].splitlines():
            ls = line.strip()
            if ls and not ls.startswith(('"', "“", "(", "[")):
                if len(ls) > 12 and not skip_stage.match(ls):
                    concept["setting"] = ls
                    break

    return concept


def _parse_video_doc(text):
    """Route for submissions using the VIDEO N: / VNA — structure."""
    today = date.today().strftime("%m%d%y")
    global_meta = _extract_global_metadata(text)
    concepts = []
    idx = 0

    # Split on VIDEO N: boundaries (keep delimiter by using lookahead)
    vid_blocks = re.split(r"(?m)^(?=VIDEO\s+\d+\s*:)", text)

    for vid_block in vid_blocks:
        if not re.match(r"VIDEO\s+\d+\s*:", vid_block.strip()):
            continue

        # Split off sub-variants (V1A —, V1B —, V3A —, etc.)
        sub_blocks = re.split(r"(?m)^(?=V\d+[A-Z]\s*[—–-])", vid_block)
        vid_meta = sub_blocks[0]  # shared header + metadata for this VIDEO

        if len(sub_blocks) > 1:
            for sub in sub_blocks[1:]:
                idx += 1
                c = _build_vid_concept(sub, idx, today, vid_meta)
                if c:
                    concepts.append(c)
        else:
            # No sub-variants — the VIDEO block itself is one concept
            idx += 1
            c = _build_vid_concept(vid_meta, idx, today, "")
            if c:
                concepts.append(c)

    # Apply global date if detected
    if global_meta.get("submission_date"):
        for c in concepts:
            c["submission_date"] = global_meta["submission_date"]

    # Recompute missing flags
    for c in concepts:
        c["missing"] = [f for f in ["duration", "setting", "inspiration"] if not c.get(f)]

    # Assign concept IDs (MMDDYY_HHMM_N) — timestamp ensures uniqueness across same-day batches
    now_hhmm = datetime.datetime.now().strftime("%H%M")
    for c in concepts:
        c["concept_id"] = f"{c.get('submission_date', date.today().strftime('%m%d%y'))}_{now_hhmm}_{c['index']}"

    return concepts


def parse_submission(text):
    # ── VIDEO N: format (e.g. SelfFinder_051126 style) ──────────────────
    if re.search(r"(?m)^VIDEO\s+\d+\s*:", text):
        return _parse_video_doc(text)

    # ── Split into concept blocks ────────────────────────────
    pattern = r"(?m)^(?:Script|Concept)\s+\d+\s*[:\.]"
    parts = re.split(pattern, text)
    headers = re.findall(pattern, text)

    concepts = []
    if len(parts) > 1:
        for i, (header, block) in enumerate(zip(headers, parts[1:])):
            c = _parse_concept_block(header + block, i + 1)
            if c:
                concepts.append(c)
    else:
        md_pattern = r"(?m)^##\s+(?:Concept|Script)\s+\d+"
        md_parts = re.split(md_pattern, text)
        md_headers = re.findall(md_pattern, text)
        if len(md_parts) > 1:
            for i, (header, block) in enumerate(zip(md_headers, md_parts[1:])):
                c = _parse_concept_block(header + block, i + 1)
                if c:
                    concepts.append(c)
        else:
            c = _parse_concept_block(text, 1)
            if c:
                concepts.append(c)

    # ── Apply trailing per-script metadata table ─────────────
    trailing = _extract_trailing_metadata(text)
    for i, meta in enumerate(trailing):
        if i >= len(concepts):
            break
        c = concepts[i]
        angle = meta.get("angle", "")
        if not c.get("duration") and meta.get("duration"):
            c["duration"] = meta["duration"]
        if not c.get("format_vibe") and meta.get("format_vibe"):
            c["format_vibe"] = meta["format_vibe"]
        # Derive setting and talent from angle if still missing
        if not c.get("setting"):
            c["setting"] = ANGLE_SETTING.get(angle, "")
        if not c.get("talent"):
            c["talent"] = ANGLE_TALENT.get(angle, "")

    # ── Apply global submission metadata ─────────────────────
    global_meta = _extract_global_metadata(text)
    if global_meta.get("submission_date"):
        for c in concepts:
            c["submission_date"] = global_meta["submission_date"]

    # ── Recompute missing flags after enrichment ─────────────
    for c in concepts:
        c["missing"] = [f for f in ["duration", "setting", "inspiration"] if not c.get(f)]

    # ── Assign concept IDs (MMDDYY_HHMM_N) ──────────────────
    # Timestamp captured once per batch so all concepts in one upload share the same HHMM
    now_hhmm = datetime.datetime.now().strftime("%H%M")
    for c in concepts:
        c["concept_id"] = f"{c.get('submission_date', date.today().strftime('%m%d%y'))}_{now_hhmm}_{c['index']}"

    return concepts
